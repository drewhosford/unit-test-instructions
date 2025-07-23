import os
import re
import yaml
import argparse
from copy import deepcopy
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

class Requirement:
    def __init__(self):
        self.req_text = ''
        self.test_steps = []
        self.test_verifications = []
        self.test_file_path = ''
        self.test_file_line = 0
        self.test_filename = ''
    
    @property
    def req_orig_text(self):
        if not hasattr(self, '_req_orig_text'):
            self._req_orig_text = ''
        return self._req_orig_text

    @req_orig_text.setter
    def req_orig_text(self, value):
        self._req_orig_text = value
        self.req_text = self.convert_camel_case_to_sentence_case(value)
        self.req_text = self.apply_conversions(self.req_text)
        self.req_text = self.apply_cleanups(self.req_text)
        first_char = self.req_text[0]
        self.req_text = first_char.upper() + self.req_text[1:]
        self.test_filename = os.path.basename(self.test_file_path)    
    
    @property
    def cleanups(self):
        return [  # the order of these items is important. If a previous item's value is a substring of a later item's key, it will be replaced
            (r'(\w)\s*/\s*(\w)',r'\1 /\2'), # convert "For a POST request to the/organizations" to "For a POST request to the /organizations"
            (r'(\w)\s*/\s*([\w-]+?)\s*/\s*(\w)',r'\1 /\2/\3'),
            (r'([A-Za-z])(\d)',r'\1 \2'),
            (r'(\d+), d,(\d+), d,(\d+), d,(\d+) /(\d+)', r'\1.\2.\3.\4/\5'),  # converts IP addresses from '192, d, 168, d, 1, d, 1 /24' to 192.168.1.1/24
            (r'([\w]), d,\s*([\w])', r'\1.\2'), # converts '192, d, 168' to 192.168
            (r'(.)" (.+?)"\s*', r'\1 "\2" '), # converts 'sets this value to" self"' to 'sets this value to "self"'
            (r' (/[\w\d]+?)-([\w\d]+?) (\w+?) ([/\w]+?) ', r' \1\2\3 '), # convert ' /exam- media /all endpoint' to '/exam-media/all endpoint'
            (r's 3', 's3'),  # convert 's 3' to 's3'
        ]

    @property
    def conversions(self):
        return [  # the order of these items is important. If a previous item's value is a substring of a later item's key, it will be replaced
            ('_ q_','"'),
            ('_ a_','&'),
            ('_ l_','<'),
            ('_ g_','>'),
            ('_ s_','/'),
            ('_ c_',':'),
            ('_ p_','|'),
            ('_ sc_',';'),
            ('_ eq_','='),
            ('_ st_','*'),
            ('_ h_','-'),
            ('_' ,','),
            ('p o s t ','POST '),
            ('g e t ','GET '),
            ('p u t ','PUT '),
            ('d e l e t e ','DELETE '),
            (' aws ',' AWS '),
            (' vpc ',' VPC '),
            (' sms',' SMS'),
            (' cidr ',' CIDR '),
            (' acl ',' ACL '),
            (' http ',' HTTP '),
            (' api ',' API '),
            (' dns ',' DNS '),
            (' tls',' TLS'),
            (' Oauth',' OAuth'),
            ('s 3 ', 'S3 '),
            ('ios ', 'iOS ')
        ]
    
    def apply_conversions(self, input_string):
        for key, value in self.conversions:
            input_string = input_string.replace(key, value)
        return input_string
    
    def apply_cleanups(self, input_string):
        for regex, value in self.cleanups:
            res = re.findall(regex, input_string)
            if len(res) > 0:
                input_string = re.sub(regex, value, input_string)
        return input_string

    # create a function that receives a camel_case string and replaces all capital letters with a space and the same letter lowercased
    def convert_camel_case_to_sentence_case(self, camel_case_string):
        result = ''
        for char in camel_case_string:
            if char.isupper():
                result += ' ' + char.lower()
            else:
                result += char
        return result.strip()        

    def __repr__(self):
        return f"Requirement: {self.req_text}\n\tTest steps: {self.test_steps}\n\tTest verifications: {self.test_verifications}"

class Section:
    def __init__(self):
        self.name = ''
        self.filenames = []
        self.requirements = []
        self.path_key = ''

class TestGroup:
    def __init__(self):
        self.file_path = ''
        self.filename = ''
        self.requirements = []

class CreateFDADocumentation:
    def __init__(self, debug_print=False, config_file_path='config.yaml'):
        self.debug_print = debug_print
        self.config_file_path = config_file_path
        pass
    
    #using yaml config file for better readability and structure
    @property
    def config(self):
        if not hasattr(self, '_config'):
            if not os.path.exists(self.config_file_path):
                print(f"Error: {self.config_file_path} file not found")
                raise Exception(f'{self.config_file_path} file not found. Please create it and put in the same directory as this script.')
            try:
                with open(self.config_file_path, 'r') as file:
                    self._config = yaml.safe_load(file)
            except Exception as e:
                raise Exception(f"Error: {self.config_file_path} file could not be parsed. {e}")
        return self._config

    def get_test_files_for_language(self, repo_path, language):
        """Recursively searches repo_path for test files by detecting testing framework imports"""
        lang_config = self.get_language_config(language)
        test_file_ext = lang_config['test_file_ext']
        test_import_patterns = lang_config['test_import_patterns']
        
        found_test_files = []
        
        # Walk through all directories in the repo
        for root, dirs, files in os.walk(repo_path):
            # Skip common non-source directories
            dirs[:] = [d for d in dirs if d not in ['.git', 'node_modules', '__pycache__', '.venv', 'venv', 'build', 'dist', 'target']]
            
            for file_name in files:
                # Check if file has the correct extension
                if not file_name.endswith(test_file_ext):
                    continue
                
                file_path = os.path.join(root, file_name)
                
                # Check if file contains testing framework imports
                if self._is_test_file(file_path, test_import_patterns):
                    found_test_files.append(file_path)
        
        return found_test_files
    
    def _is_test_file(self, file_path, test_import_patterns):
        """Check if a file contains any of the testing framework import patterns"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                # Read only the first 50 lines to check for imports (performance optimization)
                lines_to_check = 50
                for i, line in enumerate(file):
                    if i >= lines_to_check:
                        break
                    
                    # Check each test import pattern
                    for pattern in test_import_patterns:
                        if pattern.search(line):
                            return True
            return False
        except (IOError, UnicodeDecodeError):
            # If we can't read the file, assume it's not a test file
            return False
    
    def parse_file_for_requirements(self, file_path, req_regexes, test_step_re, test_verification_re, requirement_group=1):
        requirements = []
        with open(file_path, 'r') as file:
            lines = file.readlines()
            curr_req = None
            for i, line in enumerate(lines):
                # Check all requirement regex patterns
                req_match = None
                for req_re in req_regexes:
                    req_match = req_re.match(line)
                    if req_match:
                        break
                
                if req_match:
                    if curr_req:
                        requirements.append(curr_req)
                    curr_req = Requirement()
                    curr_req.test_file_path = file_path
                    curr_req.test_file_line = i + 1
                    curr_req.req_orig_text = req_match.group(requirement_group)
                    continue
                
                # Special handling for blocTest pattern where description might be on next line
                if 'blocTest<' in line and not req_match:
                    # Look for description in next few lines
                    for j in range(1, 4):  # Check next 3 lines
                        if i + j < len(lines):
                            next_line = lines[i + j]
                            desc_match = re.match(r'^\s*[\'"](.+?)[\'"],?\s*$', next_line)
                            if desc_match:
                                if curr_req:
                                    requirements.append(curr_req)
                                curr_req = Requirement()
                                curr_req.test_file_path = file_path
                                curr_req.test_file_line = i + 1
                                curr_req.req_orig_text = desc_match.group(1)
                                break
                    continue
                
                ts_match = test_step_re.match(line)
                if ts_match:
                    if not curr_req:
                        print(f"  Error: Test step found without a requirement {os.path.basename(file_path)}: {i} - '{line}'")
                    else:
                        curr_req.test_steps.append(ts_match.group(1))
                    continue
                tv_match = test_verification_re.match(line)
                if tv_match:
                    if not curr_req:
                        print(f"  Error: Test verification found without a requirement {os.path.basename(file_path)}: {i} - '{line}'")
                    else:
                        curr_req.test_verifications.append(tv_match.group(1))
                    continue
        if curr_req:
            requirements.append(curr_req)
        return requirements
    
    def get_language_config(self, language):
        """Returns language-specific configuration for parsing test files"""
        language_configs = {
            'golang': {
                'test_file_ext': '.go',
                'regex_requirement': [
                    re.compile(r'^func\s+Test_*(.+)\(t\s+\*testing.T\)'),
                ],
                'regex_test_step': re.compile(r'\s+//\s*(S\d+:\s*.+)'),
                'regex_test_ver': re.compile(r'\s+//\s*(V\d+:\s*.+)'),
                'requirement_group': 1,
                'test_import_patterns': [
                    re.compile(r'^\s*import\s+"testing"'),
                    re.compile(r'^\s*import\s+\(\s*.*"testing".*\)'),
                    re.compile(r'^\s*"testing"')
                ]
            },
            'swift': {
                'test_file_ext': '.swift',
                'regex_requirement': [
                    re.compile(r'\s+func\s+test_*(.+)\(\).+\{'),
                ],
                'regex_test_step': re.compile(r'\s+//\s*(S\d+:\s*.+)'),
                'regex_test_ver': re.compile(r'\s+//\s*(V\d+:\s*.+)'),
                'requirement_group': 1,
                'test_import_patterns': [
                    re.compile(r'^\s*import\s+XCTest'),
                    re.compile(r'^\s*@testable\s+import'),
                    re.compile(r'class\s+\w+.*:\s*XCTestCase')
                ]
            },
            'python': {
                'test_file_ext': '.py',
                'regex_requirement': [
                    re.compile(r'^\s*def\s+test_*(.+)\('),
                    re.compile(r'^\s*def\s+test(.+)\(')
                ],
                'regex_test_step': re.compile(r'\s*#\s*(S\d+:\s*.+)'),
                'regex_test_ver': re.compile(r'\s*#\s*(V\d+:\s*.+)'),
                'requirement_group': 1,
                'test_import_patterns': [
                    re.compile(r'^\s*import\s+unittest'),
                    re.compile(r'^\s*from\s+unittest'),
                    re.compile(r'^\s*import\s+pytest'),
                    re.compile(r'^\s*from\s+pytest'),
                    re.compile(r'class\s+\w+.*\(.*unittest\.TestCase.*\)')
                ]
            },
            'javascript': {
                'test_file_ext': '.js',
                'regex_requirement': [
                    re.compile(r'^\s*(it|test)\(\s*[\'"](.+?)[\'"]'),
                    re.compile(r'^\s*(it|test)\s*\(\s*[\'"](.+?)[\'"]')
                ],
                'regex_test_step': re.compile(r'\s*//\s*(S\d+:\s*.+)'),
                'regex_test_ver': re.compile(r'\s*//\s*(V\d+:\s*.+)'),
                'requirement_group': 2,
                'test_import_patterns': [
                    re.compile(r'^\s*(import|require).*[\'"]jest[\'"]'),
                    re.compile(r'^\s*(import|require).*[\'"]mocha[\'"]'),
                    re.compile(r'^\s*(import|require).*[\'"]chai[\'"]'),
                    re.compile(r'^\s*describe\s*\('),
                    re.compile(r'^\s*(it|test)\s*\(')
                ]
            },
            'typescript': {
                'test_file_ext': '.ts',
                'regex_requirement': [
                    re.compile(r'^\s*(it|test)\(\s*[\'"](.+?)[\'"]'),
                    re.compile(r'^\s*(it|test)\s*\(\s*[\'"](.+?)[\'"]')
                ],
                'regex_test_step': re.compile(r'\s*//\s*(S\d+:\s*.+)'),
                'regex_test_ver': re.compile(r'\s*//\s*(V\d+:\s*.+)'),
                'requirement_group': 2,
                'test_import_patterns': [
                    re.compile(r'^\s*import.*[\'"]jest[\'"]'),
                    re.compile(r'^\s*import.*[\'"]mocha[\'"]'),
                    re.compile(r'^\s*import.*[\'"]chai[\'"]'),
                    re.compile(r'^\s*describe\s*\('),
                    re.compile(r'^\s*(it|test)\s*\(')
                ]
            },
            'java': {
                'test_file_ext': '.java',
                'regex_requirement': [
                    re.compile(r'^\s*@DisplayName\(\s*[\'"](.+?)[\'"]\s*\)\s*'),
                    re.compile(r'^\s*@Test\s*\n\s*public\s+void\s+(.+?)\('),
                    re.compile(r'^\s*void\s+test_*(.+?)\(')
                ],
                'regex_test_step': re.compile(r'\s*//\s*(S\d+:\s*.+)'),
                'regex_test_ver': re.compile(r'\s*//\s*(V\d+:\s*.+)'),
                'requirement_group': 1,
                'test_import_patterns': [
                    re.compile(r'^\s*import\s+org\.junit'),
                    re.compile(r'^\s*import\s+org\.testng'),
                    re.compile(r'^\s*@Test'),
                    re.compile(r'^\s*@BeforeEach'),
                    re.compile(r'^\s*@AfterEach')
                ]
            },
            'csharp': {
                'test_file_ext': '.cs',
                'regex_requirement': [
                    re.compile(r'^\s*\[Test\]\s*\n\s*public\s+void\s+Test(.+?)\('),
                    re.compile(r'^\s*\[Test\]\s*public\s+void\s+Test(.+?)\('),
                    re.compile(r'^\s*\[TestMethod\]\s*\n\s*public\s+void\s+(.+?)\('),
                    re.compile(r'^\s*\[Fact\]\s*\n\s*public\s+void\s+(.+?)\(')
                ],
                'regex_test_step': re.compile(r'\s*//\s*(S\d+:\s*.+)'),
                'regex_test_ver': re.compile(r'\s*//\s*(V\d+:\s*.+)'),
                'requirement_group': 1,
                'test_import_patterns': [
                    re.compile(r'^\s*using\s+NUnit\.Framework'),
                    re.compile(r'^\s*using\s+Microsoft\.VisualStudio\.TestTools'),
                    re.compile(r'^\s*using\s+Xunit'),
                    re.compile(r'^\s*\[Test\]'),
                    re.compile(r'^\s*\[TestMethod\]'),
                    re.compile(r'^\s*\[Fact\]')
                ]
            },
            'dart': {
                'test_file_ext': '.dart',
                'regex_requirement': [
                    re.compile(r'^\s*test\(\s*[\'"](.+?)[\'"]'),
                    re.compile(r'^\s*testWidgets\(\s*[\'"](.+?)[\'"]'),
                    re.compile(r'^\s*blocTest<.+?>\s*\(\s*[\'"](.+?)[\'"]')
                ],
                'regex_test_step': re.compile(r'\s*//\s*(S\d+:\s*.+)'),
                'regex_test_ver': re.compile(r'\s*//\s*(V\d+:\s*.+)'),
                'requirement_group': 1,
                'test_import_patterns': [
                    re.compile(r'^\s*import\s+[\'"]package:test/test\.dart[\'"]'),
                    re.compile(r'^\s*import\s+[\'"]package:flutter_test/flutter_test\.dart[\'"]'),
                    re.compile(r'^\s*import\s+[\'"]package:mockito/mockito\.dart[\'"]'),
                    re.compile(r'^\s*import\s+[\'"]package:bloc_test/bloc_test\.dart[\'"]'),
                    re.compile(r'^\s*test\s*\('),
                    re.compile(r'^\s*group\s*\('),
                    re.compile(r'^\s*testWidgets\s*\('),
                    re.compile(r'^\s*blocTest\s*<')
                ]
            }
        }
        
        if language.lower() not in language_configs:
            raise Exception(f"Unsupported language: {language}. Supported languages: {', '.join(language_configs.keys())}")
        
        return language_configs[language.lower()]

    def create_documentation_from_tests(self, repo_name, repo_config):
        """Generic method to create documentation for any language based on section configuration"""
        language = repo_config.get('language', '').lower()
        if not language:
            raise Exception(f"No language specified for repo '{repo_name}'")

        # Check for required tag field
        tag = repo_config.get('tag')
        req_template_path = repo_config.get('req_template_path', '')
        ver_template_path = repo_config.get('ver_template_path', '')
        req_output_name = repo_config.get('req_output_name', '')
        ver_output_name = repo_config.get('ver_output_name', '')
        # if any of the required fields are missing, raise an exception. Allow output of all missing fields
        error_msg = []
        if not tag:
            error_msg.append(f"No tag specified for repo '{repo_name}'. Tag is required in config.")
        if not req_template_path:
            error_msg.append(f"No req_template_path specified for repo '{repo_name}'. Template is required in config.")
        if not ver_template_path:
            error_msg.append(f"No ver_template_path specified for repo '{repo_name}'. Template is required in config.")
        if not req_output_name:
            error_msg.append(f"No req_output_name specified for repo '{repo_name}'. Output name is required in config.")
        if not ver_output_name:
            error_msg.append(f"No ver_output_name specified for repo '{repo_name}'. Output name is required in config.")
        if len(error_msg) > 0:
            print("Skipping section due to missing configurations:")
            for msg in error_msg:
                print(f"  - {msg}")
            return
            
        # Get language-specific configuration
        lang_config = self.get_language_config(language)
        
        # Get sections for this documentation type
        sections = self.get_sections_for_repo(repo_name)

        # Get repository path to search for test files
        repo_path = repo_config.get('repo_path', '')
        if not repo_path:
            error_msg.append(f"No repo_path specified for repo '{repo_name}'. Repository path is required in config.")
            print("Skipping section due to missing configurations:")
            for msg in error_msg:
                print(f"  - {msg}")
            return
            
        # Find test files by searching the repository for files with testing imports
        test_files = self.get_test_files_for_language(repo_path, language)
        
        if not test_files:
            print(f"  Warning: No test files found for {language} in {repo_path}")
            return
            
        print(f"  Found {len(test_files)} test files for {language}")
        
        # Create documentation using the generic method
        self.create_documentation(
            test_files=test_files,
            lang_config=lang_config,
            sections=sections,
            tag=tag,
            template_req_doc_path=repo_config.get('req_template_path'),
            template_ver_doc_path=repo_config.get('ver_template_path'),
            output_req_doc_path=repo_config.get('req_output_name'),
            output_ver_doc_path=repo_config.get('ver_output_name')
        )
    
    def create_all_documentation(self):
        """Automatically create documentation for all configured sections"""
        for repo_name, repo_config in self.config.items():
            if not isinstance(repo_config, dict):
                continue
            
            language = repo_config.get('language', '').lower()
            
            if language:
                try:
                    print(f"Creating {language.capitalize()} documentation for section: {repo_name}")
                    self.create_documentation_from_tests(repo_name, repo_config)
                except Exception as e:
                    print(f"Error creating documentation for section '{repo_name}': {e}")
            else:
                print(f"Warning: No language specified for section '{repo_name}'. Skipping.")
        pass

    def create_documentation(self, test_files, lang_config, sections, tag, template_req_doc_path, template_ver_doc_path, output_req_doc_path, output_ver_doc_path):
        # Parse requirements from all test files
        requirements = []
        for test_file in test_files:
            print(f"Processing test file: {os.path.basename(test_file)}", end='\r')
            new_reqs = self.parse_file_for_requirements(
                test_file, 
                lang_config['regex_requirement'], 
                lang_config['regex_test_step'], 
                lang_config['regex_test_ver'], 
                lang_config['requirement_group']
            )
            requirements += new_reqs
        print(f"Found {len(requirements)} requirements from {len(test_files)} test files.")
        for req in requirements:
            section = self.get_section_for_requirement(req, sections)
            section.requirements.append(req)
        # Do some cleanup of the requirements based on certain parameters    
        for section in sections:
            if section.name == 'Ignore':
                continue
            for req in section.requirements:
                # Remove requirements that have no verification steps, and report them
                if len(req.test_verifications) == 0:
                    print(f"  Warning: No verifications in '{os.path.basename(req.test_file_path)}:{req.test_file_line}' - '{req.req_orig_text}'")
                    section.requirements.remove(req)
                # Go through the test verifications and sort them based on number in the the V\d+ tag
                req.test_verifications.sort(key=lambda x: int(x.split(':')[0][1:]))
        
        self.create_requirements_document(sections, tag, template_req_doc_path, output_req_doc_path)
        self.create_verification_document(sections, tag, template_ver_doc_path, output_ver_doc_path)

    def get_sections_for_repo(self, repo_name):
        config_repo = self.config.get(repo_name)    
        if not config_repo:
            raise Exception(f"Error: config.yaml file does not contain a section for '{repo_name}'")
        
        sections_list = config_repo.get('sections', [])
        sections = []
        
        for section_dict in sections_list:
            section = Section()
            section.name = section_dict.get('display_name', '')
            section.filenames = section_dict.get('filenames', [])
            section.path_key = section_dict.get('path_key', '')
            sections.append(section)
            
        # Add ignore section if it exists
        ignore_list = config_repo.get('ignore', [])
        if ignore_list:
            ignore_section = Section()
            ignore_section.name = "Ignore"
            ignore_section.filenames = ignore_list
            sections.append(ignore_section)
            
        # Add miscellaneous section for unmatched files
        misc_section = Section()
        misc_section.name = "Miscellaneous"
        sections.append(misc_section)
        return sections
    
    def get_section_for_requirement(self, req, sections):
        misc_section = None
        for section in sections:
            if section.path_key != '' and section.path_key not in req.test_file_path:
                continue
            if req.test_filename in section.filenames:
                return section
            if section.name == 'Miscellaneous':
                misc_section = section
        return misc_section

    def get_tag_style(self, document):
        style2 = document.styles['Heading 2']
        try:
            style = document.styles['tag_font']
        except KeyError:
            style = None
        if not style:
            style = document.styles.add_style('tag_font', WD_STYLE_TYPE.CHARACTER)
            style.font.size = style2.font.size
            style.font.name = style2.font.name
            style.font.color.rgb = RGBColor(255, 0, 0)
            style.bold = True
        return style

    def create_requirements_document(self, sections, tag, docx_path, output_docx_name):
        document = Document(docx_path)
        style2 = document.styles['Heading 2']
        tag_style = self.get_tag_style(document)
        # for each section in the dictionary, create a new section in the document with a style of Heading 1, and each requirement in that section with a style of Heading 2
        req_num = 0
        for section in sections:
            if section.name == 'Ignore':
                continue
            if len(section.requirements) == 0:
                continue
            document.add_heading(section.name, level=1)
            for req in section.requirements:
                req_num += 1
                req.req_num = req_num
                p = document.add_paragraph()
                p.style = style2
                run1 = p.add_run(f"{tag}:DO:{req_num}")
                p.add_run(f" {req.req_text}")
                if self.debug_print:
                    p = p.add_run(f" ({os.path.basename(req.test_file_path)}:{req.test_file_line})")
                    p.italic = True
                run1.style = tag_style
                run1.bold = True
        
        # Create Outputs directory if it doesn't exist
        dirname = os.path.dirname(docx_path)
        outputs_dir = os.path.join(dirname, 'Outputs')
        if not os.path.exists(outputs_dir):
            os.makedirs(outputs_dir)
            print(f"Created Outputs directory: {outputs_dir}")
        
        filename = os.path.join(outputs_dir, output_docx_name)
        document.save(filename)
        print(f"Saved requirements document: {filename}")
        pass

    def create_verification_document(self, sections, tag, docx_path, output_docx_name):
        document = Document(docx_path)
        # Get the paragraph with the text "Test Steps", if the style does not exist, create it
        test_step_num_style = document.styles['Normal']
        tag_style = self.get_tag_style(document)
        p = document.add_paragraph()
        p.style = document.styles['Heading 1']
        p.add_run("Verification Test Protocol")
        ver_num = 1
        ver_step_num = 1
        # get the last table
        table = document.tables[-1]
        # make a deep copy of the table
        copied_table_tbl = deepcopy(table._tbl)
        # Find the paragraph just before the last table
        first_section_paragraph = None
        for element in document.element.body:
            if element.tag.endswith('tbl') and element == table._tbl:
                # Found the table, get the previous element
                prev_element = element.getprevious()
                if prev_element is not None and prev_element.tag.endswith('p'):
                    # Find the corresponding paragraph object
                    for para in document.paragraphs:
                        if para._element == prev_element:
                            first_section_paragraph = para
                            break
                break
        
        for table_i, section in enumerate(sections):
            if section.name == 'Ignore':
                continue
            if len(section.requirements) == 0:
                print(f"  Warning: No requirements in section '{section.name}' for tag '{tag}'. Skipping section.")
                continue
            
            if table_i == 0 and first_section_paragraph is not None:
                # Use the existing paragraph before the table for the first section
                p = first_section_paragraph
                # Clear existing content and add new content
                p.clear()
            else:
                p = document.add_paragraph()
            p.style = document.styles['Heading 2']
            run = p.add_run(f"{tag}:VER:{ver_num} ")
            run.style = tag_style
            run.bold = True
            p.add_run(section.name)
            if table_i > 0:
                # Create a new table from the copied XML element and insert it
                new_table_tbl = deepcopy(copied_table_tbl)
                p._p.addnext(new_table_tbl)
                # Get the newly added table from the document's tables collection
                table = document.tables[-1]  # The newly added table will be the last one
            # ensure that the table has sufficient rows for all the requirements (keeping in mind that the copied table already has a header row and an empty first row)
            for i in range(len(section.requirements) - 1):
                table.add_row()
            for i, req in enumerate(section.requirements):
                cells = table.rows[i + 1].cells
                cells[0].text = f"{ver_step_num}."
                ver_step_num += 1
                cells[0].paragraphs[0].style = test_step_num_style
                cells[1].text = '\n'.join(req.test_steps)
                cells[2].text = '\n'.join(req.test_verifications)
                p = cells[2].add_paragraph()
                run = p.add_run(f"[{tag}:DO:{req.req_num}]")
                run.style = tag_style
                run.bold = True
                if self.debug_print:
                    p = cells[2].add_paragraph()
                    p.text = f"({os.path.basename(req.test_file_path)}:{req.test_file_line})"
                    p.italic = True
            ver_num += 1
        
        # Create Outputs directory if it doesn't exist
        dirname = os.path.dirname(docx_path)
        outputs_dir = os.path.join(dirname, 'Outputs')
        if not os.path.exists(outputs_dir):
            os.makedirs(outputs_dir)
            print(f"Created Outputs directory: {outputs_dir}")
        
        filename = os.path.join(outputs_dir, output_docx_name)
        document.save(filename)
        print(f"Saved verification document: {filename}")
        pass


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description='Create FDA documentation')
    parser.add_argument('--config', '-c', type=str, default='config.yaml', 
                       help='Path to the configuration file (default: config.yaml)')
    args = parser.parse_args()
    # Create an instance of the CreateFDADocumentation class with the specified config file
    create_fda_documentation = CreateFDADocumentation(debug_print=False, config_file_path=args.config)
    
    print(f"Processing all configured sections from {args.config}...")
    create_fda_documentation.create_all_documentation()





